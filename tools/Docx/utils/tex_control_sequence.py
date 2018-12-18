"""
TeX Control Sequence.
"""


from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font

import re


# Constant(s)
EXTENSION = ".docx"

# Function(s)
def begindocument():
    r"""
    \begin{document}
    -------------
    Arg(s):
        None.

    Return(s):
        Document.
    """
    return Document()


def enddocument(document, file:str="demo.docx"):
    r"""
    \end{document}
    -------------
    Arg(s):
        Document.

    Return(s):
        None.
    """
    document.save(file + ("" if file.endswith(EXTENSION) else EXTENSION))


def title(document, title:str):
    r"""
    \title{title}.
    -------------
    Arg(s):
        document: Document.
        title: String.

    Return(s):
        None.
    """
    document.add_heading(title, level=0)


def section(document, name:str):
    r"""
    \section{name}.
    -------------
    Arg(s):
        document: Document.
        name: String.

    Return(s):
        None.
    """
    document.add_heading(name, level=1)


def subsection(document, name:str):
    r"""
    \subsection{name}.
    -------------
    Arg(s):
        document: Document.
        name: String.

    Return(s):
        None.
    """
    document.add_heading(name, level=2)


def subsubsection(document, name:str):
    r"""
    \subsubsection{name}.
    -------------
    Arg(s):
        document: Document.
        name: String.

    Return(s):
        None.
    """
    document.add_heading(name, level=3)


def includegraphics(document, image:str, opt:dict=None):
    r"""
    \includegraphics[opt]{image}.
    -------------
    Arg(s):
        document: Document.
        image: String.
        opt: Dictionary, {"width":Inches(1), "height":Inches(1)}.

    Return(s):
        None.
    """
    document.add_picture(image, **opt)


def paragraph(document, content:str):
    r"""
    Paragraph.
    -------------
    Arg(s):
        document: Document.
        content: String.

    Return(s):
        None.
    """
    ITALIC    = ["*",   "*"]
    BOLD      = ["**",  "**"]
    BOLDIT    = ["***", "***"]
    UNDERLINE = ["<u>", "</u>"]
    stack     = []
    result    = ""
    document.add_paragraph(content)
    # TODO, convert content using markdown grammar.
    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic').italic = True
    # p.add_run(' and some ')
    # tmp = p.add_run('underline.')
    # tmp.underline = True
    # tmp.italic = True
    # tmp.bold = True


def centering(document):
    r"""
    \centering{document}.
    -------------
    Arg(s):
        document: Document.

    Return(s):
        None.
    """
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def flushleft(document):
    r"""
    \flushleft{document}.
    -------------
    Arg(s):
        document: Document.

    Return(s):
        None.
    """
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT


def flushright(document):
    r"""
    \flushleft{document}.
    -------------
    Arg(s):
        document: Document.

    Return(s):
        None.
    """
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def newpage(document):
    r"""
    \newpage.
    -------------
    Arg(s):
        document: Document.

    Return(s):
        None.
    """
    document.add_page_break()




document = begindocument()
title(document, "TITLE")
section(document, "section")
subsection(document, "subsection")
subsubsection(document, "subsubsection")
includegraphics(document, "400x400.bmp", {"width":Inches(1)})
centering(document)
includegraphics(document, "400x400.bmp", {"width":Inches(1)})
flushleft(document)
includegraphics(document, "400x400.bmp", {"width":Inches(1)})
flushright(document)
section(document, "Paragraph.")
content  = "This is *italic*, this is **bold**, this is <u>underline</u>.\n"
content += "This is <u>***ALL***</u>.\n"
content += "*Test **all** <u>feature</u>*."
paragraph(document, content)
enddocument(document, "test")
